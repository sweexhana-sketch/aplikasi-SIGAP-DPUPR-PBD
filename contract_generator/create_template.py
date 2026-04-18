
from docx import Document
import os

target_file = "SURAT PERJANJIAN Pekerjaan Konstruksi Harga satuan 2021.docx"
output_file = "Template_Master_2024.docx"

replacements = {
    # 1. Judul Paket (Contoh placeholder yang mungkin ada di file asli)
    "Peningkatan Jalan diPetrocina/Jalan Lingkar": "{{ nama_paket }}",
    "[Nama Paket Pekerjaan]": "{{ nama_paket }}", 
    
    # 2. Nomor Kontrak
    "Nomor : ........................ [diisi nomor Kontrak]": "Nomor : {{ no_kontrak }}",
    
    # 3. Tanggal
    "pada hari .......... tanggal ….... bulan ...": "pada tanggal {{ tgl_kontrak }}",
    
    # 4. Nama Penyedia
    "[Nama Penyedia]": "{{ nama_penyedia }}",
    
    # 5. Waktu
    "selama ………. (… dalam huruf …) hari kalender": "selama {{ waktu }} hari kalender",
    
    # 6. Nilai Kontrak
    "Rp .................": "Rp {{ nilai_kontrak }}",
    
    # 7. Terbilang
    "(…………………………..)": "({{ terbilang }})",
    
    # 8. SPMK
    "Nomor SPMK: ........................": "Nomor SPMK: {{ no_spmk }}"
}

# Fungsi Helper untuk replace text dalam paragraph tanpa merusak format run
def replace_text_in_paragraph(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        # Simple string replacement (Note: This might break complex formatting if 'old_text' is split across runs)
        # But for template markers it is usually safe enough. 
        # For rigorous replacement we need to check runs, but python-docx is tricky with that.
        # We will try a naive approach first which works 80% of time for simple documents.
        paragraph.text = paragraph.text.replace(old_text, new_text)

def main():
    if not os.path.exists(target_file):
        print(f"Error: File '{target_file}' tidak ditemukan.")
        return

    doc = Document(target_file)
    print(f"Memproses file: {target_file}...")

    # Iterate through paragraphs
    total_replaced = 0
    for para in doc.paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                replace_text_in_paragraph(para, old, new)
                print(f"Replaced in Paragraph: '{old}' -> '{new}'")
                total_replaced += 1

    # Iterate through tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in para.text:
                            replace_text_in_paragraph(para, old, new)
                            print(f"Replaced in Table: '{old}' -> '{new}'")
                            total_replaced += 1
    
    doc.save(output_file)
    print(f"\nSelesai! File template disimpan sebagai: {output_file}")
    print(f"Total penggantian dilakukan: {total_replaced}")

if __name__ == "__main__":
    main()
